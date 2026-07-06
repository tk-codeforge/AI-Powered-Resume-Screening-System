"""
resume_parser.py - Extracts text from PDF and DOCX resumes
"""
import re
import pdfplumber
from docx import Document
from pathlib import Path


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file using pdfplumber."""
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {str(e)}")
    return text.strip()


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file."""
    text = ""
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
            text += "\n"
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}")
    return text.strip()


def extract_text(file_path: str) -> str:
    """Auto-detect file type and extract text."""
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return extract_text_from_pdf(file_path)
    elif suffix in [".docx", ".doc"]:
        return extract_text_from_docx(file_path)
    elif suffix == ".txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


def extract_email(text: str) -> str | None:
    """Extract email address from text."""
    pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    match = re.search(pattern, text)
    return match.group(0) if match else None


def extract_phone(text: str) -> str | None:
    """Extract phone number from text."""
    pattern = r'(\+?\d{1,3}[\s\-.]?)?\(?\d{3}\)?[\s\-.]?\d{3}[\s\-.]?\d{4}'
    match = re.search(pattern, text)
    return match.group(0) if match else None


def extract_name(text: str) -> str | None:
    """Extract candidate name (first non-empty line, heuristic)."""
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    if lines:
        first_line = lines[0]
        # Likely a name if it's 2-4 words and no special chars
        words = first_line.split()
        if 2 <= len(words) <= 4 and all(w.replace('-', '').isalpha() for w in words):
            return first_line
    return None


def parse_resume(file_path: str) -> dict:
    """
    Parse a resume file and return structured data.
    Returns: { name, email, phone, raw_text }
    """
    raw_text = extract_text(file_path)
    return {
        "name": extract_name(raw_text),
        "email": extract_email(raw_text),
        "phone": extract_phone(raw_text),
        "raw_text": raw_text,
        "word_count": len(raw_text.split()),
        "char_count": len(raw_text),
    }
